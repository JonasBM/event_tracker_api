import re
import os
import unicodedata
from datetime import date, timedelta

from docx import Document


def getDateFromString(stringDate):
    if stringDate:
        date_re = (
            "((?:19|20)\\d\\d)-(0?[1-9]|1[012])-([12][0-9]|3[01]|0?[1-9])"
        )
        matchObj = re.fullmatch(date_re, stringDate)
        if matchObj:
            split = stringDate.split("-")
            return date(int(split[0]), int(split[1]), int(split[2]))
    return None


def strip_accents(text):
    """
    Strip accents from input String.

    :param text: The input string.
    :type text: String.

    :returns: The processed String.
    :rtype: String.
    """
    try:
        text = unicode(text, "utf-8")
    except (TypeError, NameError):  # unicode is a default on python 3
        pass
    text = unicodedata.normalize("NFD", text)
    text = text.encode("ascii", "ignore")
    text = text.decode("utf-8")
    return str(text)


def text_to_id(text):
    """
    Convert input text to id.

    :param text: The input string.
    :type text: String.

    :returns: The processed String.
    :rtype: String.
    """
    text = strip_accents(text.lower())
    text = re.sub("[ ]+", "_", text)
    text = re.sub("[^0-9a-zA-Z_-]", "", text)
    return text


def add_days(from_date, number_of_days, business_days=False):
    to_date = from_date
    while number_of_days:
        to_date += timedelta(1)
        if to_date.weekday() < 5 or not business_days:
            number_of_days -= 1
    return to_date


def count_days(from_date, to_date, business_days=False):
    number_of_days = 0
    while from_date < to_date:
        from_date += timedelta(1)
        if from_date.weekday() < 5 or not business_days:
            number_of_days += 1
    return number_of_days


def add_month(data):
    month = data.month + 1
    year = data.year
    if data.month == 12:
        month = 1
        year = data.year + 1
    return date(year, month, data.day)


def sub_month(data):
    month = data.month - 1
    year = data.year
    if data.month == 1:
        month = 12
        year = data.year - 1
    return date(year, month, data.day)


def docxFromTemplate(file_path, context):
    if os.path.exists(file_path):
        print(context)
        document = Document(file_path)

        for p in document.paragraphs:
            for r in p.runs:
                for key, value in context.items():
                    if key in r.text:
                        if value:
                            r.text = r.text.replace(key, value)
                        else:
                            r.text = r.text.replace(key, "")

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            for key, value in context.items():
                                if key in r.text:
                                    if value:
                                        r.text = r.text.replace(key, value)
                                    else:
                                        r.text = r.text.replace(key, " ")

        return document
    return None
