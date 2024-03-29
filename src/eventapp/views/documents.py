import calendar
import csv
import locale
import os
from datetime import date

from django.conf import settings
from django.contrib.auth.models import User
from django.db.models import Q
from django.http import Http404, HttpResponse
from django.template.loader import get_template
from django.utils import timezone
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from eventapp.models import (
    Notice,
    NoticeEvent,
    NoticeEventType,
    NoticeEventTypeFile,
    ReportEvent,
    ReportEventType,
    SurveyEvent,
    SurveyEventType,
)
from eventapp.utils import docxFromTemplate, getDateFromString, parse_bool
from rest_framework import generics, permissions, serializers, status
from rest_framework.response import Response
from xhtml2pdf import pisa


def render_pdf_from_html(
    template_src, context_dict={}, file_name="document.pdf"
):
    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = "attachment; filename=" + file_name
    template = get_template(template_src)
    html = template.render(context_dict)
    pisa_status = pisa.CreatePDF(html, dest=response)
    if pisa_status.err:
        return HttpResponse("We had some errors <pre>" + html + "</pre>")
    return response


class ReportPDF(generics.RetrieveAPIView):
    permission_classes = [
        permissions.IsAuthenticated,
    ]

    def get(self, request, *args, **kwargs):

        user = self.request.user
        user_id = self.request.query_params.get("user_id", None)
        if user_id:
            user_instance = User.objects.filter(id=user_id).first()
            if user_instance:
                user = user_instance

        if not self.request.user.profile.is_auditor():
            if (
                not user.profile.has_my_permission(request.user)
                or self.request.user.profile.is_particular()
            ):
                return Response(status=status.HTTP_403_FORBIDDEN)

        include_analytic_data = self.request.query_params.get(
            "include_analytic_data", None
        )
        if not include_analytic_data:
            include_analytic_data = True
        if type(include_analytic_data) is not bool:
            if include_analytic_data == "false":
                include_analytic_data = False
            else:
                include_analytic_data = True

        month = self.request.query_params.get("month", None)
        if month:
            start_date = getDateFromString(month + "-01")
            end_date = getDateFromString(
                month
                + "-"
                + str(
                    calendar.monthrange(start_date.year, start_date.month)[1]
                )
            )
        query_range_notice_event = Q(
            notice_events__date__range=[start_date, end_date]
        )
        query_owner_notice = Q(notice__owner=user)
        query_range = Q(date__range=[start_date, end_date])
        query_owner = Q(owner=user)

        notice_event_types = (
            NoticeEventType.objects.distinct()
            .filter(
                query_range_notice_event & Q(notice_events__notice__owner=user)
            )
            .all()
        )

        notices = (
            Notice.objects.distinct()
            .filter(query_range_notice_event & query_owner)
            .all()
        )
        notice_events = NoticeEvent.objects.filter(
            query_range & query_owner_notice
        ).all()

        survey_event_types = (
            SurveyEventType.objects.distinct()
            .filter(
                Q(survey_events__date__range=[start_date, end_date])
                & Q(survey_events__owner=user)
            )
            .all()
        )
        survey_events = user.surveys.filter(query_range & query_owner).all()

        report_event_types = (
            ReportEventType.objects.distinct()
            .filter(
                Q(report_events__date__range=[start_date, end_date])
                & Q(report_events__owner=user)
            )
            .all()
        )
        report_events = user.reports.filter(query_range & query_owner).all()

        activitys = user.activitys.filter(query_range & query_owner).all()

        context = {
            "user": user,
            "include_analytic_data": include_analytic_data,
            "today": date.today(),
            "reference": start_date,
            "notice_event_types": notice_event_types,
            "survey_event_types": survey_event_types,
            "report_event_types": report_event_types,
            "notices": notices,
            "notice_events": notice_events,
            "survey_events": survey_events,
            "report_events": report_events,
            "activitys": activitys,
        }
        pdf = render_pdf_from_html(
            "pdf/relatorio.html", context, "relatorio.pdf"
        )
        return HttpResponse(pdf, content_type="application/pdf")


class ReportPDFAll(generics.RetrieveAPIView):
    permission_classes = [
        permissions.IsAuthenticated,
    ]

    def get(self, request, *args, **kwargs):

        if (
            self.request.user.profile.is_assistente()
            or self.request.user.profile.is_particular()
        ):
            return Response(status=status.HTTP_403_FORBIDDEN)

        month = self.request.query_params.get("month", None)
        if month:
            start_date = getDateFromString(month + "-01")
            end_date = getDateFromString(
                month
                + "-"
                + str(
                    calendar.monthrange(start_date.year, start_date.month)[1]
                )
            )
        query_range_notice_event = Q(
            notice_events__date__range=[start_date, end_date]
        )
        query_owner_notice = Q(notice__owner__is_superuser=False)
        query_range = Q(date__range=[start_date, end_date])
        query_owner = Q(owner__is_superuser=False)

        notice_event_types = (
            NoticeEventType.objects.distinct()
            .filter(
                query_range_notice_event
                & Q(notice_events__notice__owner__is_superuser=False)
            )
            .all()
        )

        notices = (
            Notice.objects.distinct()
            .filter(query_range_notice_event & query_owner)
            .all()
        )
        notice_events = NoticeEvent.objects.filter(
            query_range & query_owner_notice
        ).all()

        survey_event_types = (
            SurveyEventType.objects.distinct()
            .filter(
                Q(survey_events__date__range=[start_date, end_date])
                & Q(survey_events__owner__is_superuser=False)
            )
            .all()
        )
        survey_events = SurveyEvent.objects.filter(
            query_range & query_owner
        ).all()

        report_event_types = (
            ReportEventType.objects.distinct()
            .filter(
                Q(report_events__date__range=[start_date, end_date])
                & Q(report_events__owner__is_superuser=False)
            )
            .all()
        )
        report_events = ReportEvent.objects.filter(
            query_range & query_owner
        ).all()

        context = {
            "today": date.today(),
            "reference": start_date,
            "notice_event_types": notice_event_types,
            "survey_event_types": survey_event_types,
            "report_event_types": report_event_types,
            "notices": notices,
            "notice_events": notice_events,
            "survey_events": survey_events,
            "report_events": report_events,
        }
        pdf = render_pdf_from_html(
            "pdf/relatorio_geral.html", context, "relatorio_geral.pdf"
        )
        return HttpResponse(pdf, content_type="application/pdf")


class sheetCSV(generics.RetrieveAPIView):
    permission_classes = [
        permissions.IsAuthenticated,
    ]

    def get(self, request, *args, **kwargs):
        user = self.request.user
        user_id = self.request.query_params.get("user_id", None)
        if user_id:
            user_instance = User.objects.filter(id=user_id).first()
            if user_instance:
                user = user_instance

        if not self.request.user.profile.is_auditor():
            if (
                not user.profile.has_my_permission(request.user)
                or self.request.user.profile.is_particular()
            ):
                return Response(status=status.HTTP_403_FORBIDDEN)

        month = self.request.query_params.get("month", None)
        if month:
            start_date = getDateFromString(month + "-01")
            end_date = getDateFromString(
                month
                + "-"
                + str(
                    calendar.monthrange(start_date.year, start_date.month)[1]
                )
            )

        query_range = Q(date__range=[start_date, end_date])
        query_owner = Q(owner=user)

        query_notice = Q(notice_events__date__range=[start_date, end_date])
        query_notice_deadline = Q(
            notice_events__deadline_date__range=[start_date, end_date]
        )
        notices = user.notices.filter(
            query_notice | query_notice_deadline
        ).all()

        survey_events = user.surveys.filter(query_range & query_owner).all()
        activitys = user.activitys.filter(query_range & query_owner).all()

        filename = (
            "planilha_mensal-"
            + str(start_date.year)
            + "-"
            + str(start_date.month)
        )
        response = HttpResponse(content_type="text/csv")
        response["Content-Disposition"] = (
            'attachment; filename="' + filename + '.csv"'
        )

        response.write(u"\ufeff".encode("utf8"))

        writer = csv.writer(response, dialect=csv.excel, delimiter=";")
        writer.writerow(["AUTOS"])
        writer.writerow(["GRUPOS DE AUTOS"])
        header = ["ID", "Imóvel", "Endereço", "Descrição"]
        writer.writerow(header)
        for notice in notices:
            notice_line = [
                notice.id,
                notice.imovel.name_string,
                notice.address,
                notice.description,
            ]
            writer.writerow(notice_line)
        writer.writerow("")
        writer.writerow(["AUTOS SEPARADOS"])
        header = [
            "date",
            "Grupo do Auto (ID)",
            "Imóvel",
            "Tipo",
            "Nº",
            "RF",
            "Prazo (dias)",
            "Conta dias úteis",
            "Prazo(data)",
            "Concluido",
        ]
        writer.writerow(header)
        for notice in notices:
            for notice_event in notice.notice_events.all():
                notice_event_line = [
                    notice_event.date,
                    notice_event.notice.id,
                    notice_event.notice.imovel.name_string,
                    notice_event.notice_event_type.name,
                    notice_event.identification,
                    notice_event.report_number,
                    notice_event.deadline,
                    notice_event.deadline_working_days,
                    notice_event.deadline_date,
                    notice_event.concluded,
                ]
                writer.writerow(notice_event_line)
        writer.writerow("")
        writer.writerow("")
        writer.writerow(["VISTORIAS"])
        header = ["Data", "Tipo", "Nº", "Endereço", "Descrição", "Concluido"]
        writer.writerow(header)
        for survey in survey_events:
            survey_line = [
                survey.date,
                survey.survey_event_type.name,
                survey.identification,
                survey.address,
                survey.description,
                survey.concluded,
            ]
            writer.writerow(survey_line)
        writer.writerow("")
        writer.writerow(["ATIVIDADES"])
        header = ["Data", "Descrição"]
        writer.writerow(header)
        for activity in activitys:
            activity_line = [
                activity.date,
                activity.description.replace("\n", "\r\n"),
            ]
            writer.writerow(activity_line)
        return response


class NoticeReportDocx(generics.RetrieveAPIView):
    permission_classes = [
        permissions.IsAuthenticated,
    ]

    def get(self, request, *args, **kwargs):
        locale.setlocale(locale.LC_TIME, "pt_BR")

        notice_id = self.request.query_params.get("notice_id", None)
        if notice_id:
            notice = Notice.objects.filter(id=notice_id).first()

            if notice:

                if not self.request.user.profile.is_auditor():
                    if (
                        not notice.owner.profile.has_my_permission(request.user)
                        or self.request.user.profile.is_particular()
                    ):
                        return Response(status=status.HTTP_403_FORBIDDEN)

                if not notice.imovel:
                    return Response(
                        {
                            "detail": "Apenas é possivel gerar relatório com imóvel cadastrado"
                        },
                        status=status.HTTP_400_BAD_REQUEST,
                    )

                first_notice_event = (
                    notice.notice_events.order_by("-date").all().first()
                )
                file_path = (
                    settings.MEDIA_ROOT + "//relatorio_padrao//rf_padrao.docx"
                )
                vistoria_administrativa = notice.notice_events.filter(
                    notice_event_type__name="Vistoria Administrativa"
                ).first()

                report_number = "XXX/"
                report_number += timezone.localtime(timezone.now()).strftime(
                    "%Y"
                )

                va_identification = "YYY"

                if vistoria_administrativa:
                    if vistoria_administrativa.report_number:
                        report_number = vistoria_administrativa.report_number

                    if vistoria_administrativa.identification:
                        va_identification = (
                            vistoria_administrativa.identification
                        )

                if os.path.exists(file_path):
                    document = Document(file_path)
                    for p in document.paragraphs:
                        for r in p.runs:

                            if "data_atual_por_extenso" in r.text:
                                r.text = r.text.replace(
                                    "data_atual_por_extenso",
                                    timezone.localtime(
                                        timezone.now()
                                    ).strftime("%d de %B de %Y"),
                                )

                            if first_notice_event:
                                if "data_vistoria" in r.text:
                                    r.text = r.text.replace(
                                        "data_vistoria",
                                        first_notice_event.date.strftime(
                                            "%d de %B de %Y"
                                        ),
                                    )

                            if "numero_relatorio_fiscalizacao" in r.text:
                                r.text = r.text.replace(
                                    "numero_relatorio_fiscalizacao",
                                    report_number,
                                )

                            if "numero_VA" in r.text:
                                r.text = r.text.replace(
                                    "numero_VA",
                                    va_identification,
                                )

                            if "afm_nome_completo" in r.text:
                                r.text = r.text.replace(
                                    "afm_nome_completo",
                                    notice.owner.get_full_name(),
                                )

                            if "afm_matricula" in r.text:
                                r.text = r.text.replace(
                                    "afm_matricula",
                                    notice.owner.profile.matricula,
                                )

                            if "endereço_completo" in r.text:
                                address_string = ""
                                if notice.imovel.logradouro:
                                    address_string = notice.imovel.logradouro
                                if notice.imovel.numero:
                                    address_string += (
                                        ", n" + notice.imovel.numero
                                    )
                                if notice.imovel.complemento:
                                    address_string += (
                                        ", " + notice.imovel.complemento
                                    )
                                if notice.imovel.bairro:
                                    address_string += (
                                        " - " + notice.imovel.bairro
                                    )
                                r.text = r.text.replace(
                                    "endereço_completo",
                                    "endereço: " + address_string,
                                )

                            if "lista_de_autos" in r.text:
                                lista_autos_string = "Autos de"

                                started_list = False
                                intimacoes = notice.notice_events.filter(
                                    notice_event_type__name="Intimação"
                                ).all()
                                infracoes = notice.notice_events.filter(
                                    notice_event_type__name="Infração"
                                ).all()
                                embargos = notice.notice_events.filter(
                                    notice_event_type__name="Embargo"
                                ).all()

                                if intimacoes.count() > 0:
                                    lista_autos_string += " Intimação"
                                    started = False
                                    for intimacao in intimacoes:
                                        if started:
                                            lista_autos_string += ","
                                        lista_autos_string += (
                                            " N°" + intimacao.identification
                                        )
                                        started = True
                                    started_list = True

                                if infracoes.count() > 0:
                                    if started_list:
                                        lista_autos_string += ";"
                                    if embargos.count() == 0:
                                        lista_autos_string += " e"
                                    lista_autos_string += " Infração"
                                    started = False
                                    for infracao in infracoes:
                                        if started:
                                            lista_autos_string += ","
                                        lista_autos_string += (
                                            " N°" + infracao.identification
                                        )
                                        started = True
                                    started_list = True

                                if embargos.count() > 0:
                                    if started_list:
                                        lista_autos_string += "; e"
                                    lista_autos_string += " Embargo"
                                    started = False
                                    for embargo in embargos:
                                        if started:
                                            lista_autos_string += ","
                                        lista_autos_string += (
                                            " N°" + embargo.identification
                                        )
                                        started = True

                                r.text = r.text.replace(
                                    "lista_de_autos", lista_autos_string
                                )

                            if "numero_cadastro" in r.text:
                                r.text = r.text.replace(
                                    "numero_cadastro",
                                    "Cadastro N°" + notice.imovel.codigo,
                                )

                            if "razão_social_do_imovel" in r.text:
                                r.text = r.text.replace(
                                    "razão_social_do_imovel",
                                    notice.imovel.razao_social,
                                )

                else:
                    document = Document()
                    document.add_heading("Document Title", 0)

                response = HttpResponse(
                    content_type=(
                        "application/"
                        "vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                )
                response[
                    "Content-Disposition"
                ] = "attachment; filename=download.docx"
                document.save(response)

                return response

        return Response(status=status.HTTP_400_BAD_REQUEST)


class VARequestDocx(generics.RetrieveAPIView):
    permission_classes = [
        permissions.IsAuthenticated,
    ]

    def get(self, request, *args, **kwargs):
        locale.setlocale(locale.LC_TIME, "pt_BR")

        vistoria_administrativa_id = self.request.query_params.get(
            "vistoria_administrativa_id", None
        )
        if vistoria_administrativa_id:
            vistoria_administrativa = NoticeEvent.objects.filter(
                id=vistoria_administrativa_id
            ).first()

            if vistoria_administrativa:
                notice = vistoria_administrativa.notice

                if not self.request.user.profile.is_auditor():
                    if (
                        not notice.owner.profile.has_my_permission(request.user)
                        or self.request.user.profile.is_particular()
                    ):
                        return Response(status=status.HTTP_403_FORBIDDEN)

                if not notice.imovel:
                    return Response(
                        {
                            "detail": "Apenas é possivel gerar relatório com imóvel cadastrado"
                        },
                        status=status.HTTP_400_BAD_REQUEST,
                    )

                first_notice_event = (
                    notice.notice_events.order_by("-date").all().first()
                )
                file_path = os.path.join(
                    settings.MEDIA_ROOT, "relatorio_padrao", "va_padrao.docx"
                )
                report_number = "XXX/"
                report_number += timezone.localtime(timezone.now()).strftime(
                    "%Y"
                )
                if vistoria_administrativa:
                    if vistoria_administrativa.report_number:
                        report_number = vistoria_administrativa.report_number

                if os.path.exists(file_path):
                    document = Document(file_path)
                    for p in document.paragraphs:
                        for r in p.runs:

                            if "data_atual_por_extenso" in r.text:
                                r.text = r.text.replace(
                                    "data_atual_por_extenso",
                                    timezone.localtime(
                                        timezone.now()
                                    ).strftime("%d de %B de %Y"),
                                )

                            if "AFM_nome_completo" in r.text:
                                r.text = r.text.replace(
                                    "AFM_nome_completo",
                                    "AFM " + notice.owner.get_full_name(),
                                )

                            if "endereço_completo" in r.text:
                                address_string = ""
                                if notice.imovel.logradouro:
                                    address_string = notice.imovel.logradouro
                                if notice.imovel.numero:
                                    address_string += (
                                        ", n" + notice.imovel.numero
                                    )
                                if notice.imovel.complemento:
                                    address_string += (
                                        ", " + notice.imovel.complemento
                                    )
                                if notice.imovel.bairro:
                                    address_string += (
                                        " - " + notice.imovel.bairro
                                    )
                                r.text = r.text.replace(
                                    "endereço_completo",
                                    "endereço: " + address_string,
                                )

                            # "n° " + report_number,
                            if "numero_VA" in r.text:
                                r.text = r.text.replace(
                                    "numero_VA",
                                    "n° "
                                    + vistoria_administrativa.identification,
                                )

                            if "lista_de_autos" in r.text:
                                lista_autos_string = "Autos de"

                                started_list = False
                                intimacoes = notice.notice_events.filter(
                                    notice_event_type__name="Intimação"
                                ).all()
                                infracoes = notice.notice_events.filter(
                                    notice_event_type__name="Infração"
                                ).all()
                                embargos = notice.notice_events.filter(
                                    notice_event_type__name="Embargo"
                                ).all()

                                if intimacoes.count() > 0:
                                    lista_autos_string += " Intimação"
                                    started = False
                                    for intimacao in intimacoes:
                                        if started:
                                            lista_autos_string += ","
                                        lista_autos_string += (
                                            " N°" + intimacao.identification
                                        )
                                        started = True
                                    started_list = True

                                if infracoes.count() > 0:
                                    if started_list:
                                        lista_autos_string += ";"
                                    if embargos.count() == 0:
                                        lista_autos_string += " e"
                                    lista_autos_string += " Infração"
                                    started = False
                                    for infracao in infracoes:
                                        if started:
                                            lista_autos_string += ","
                                        lista_autos_string += (
                                            " N°" + infracao.identification
                                        )
                                        started = True
                                    started_list = True

                                if embargos.count() > 0:
                                    if started_list:
                                        lista_autos_string += "; e"
                                    lista_autos_string += " Embargo"
                                    started = False
                                    for embargo in embargos:
                                        if started:
                                            lista_autos_string += ","
                                        lista_autos_string += (
                                            " N°" + embargo.identification
                                        )
                                        started = True

                                r.text = r.text.replace(
                                    "lista_de_autos", lista_autos_string
                                )

                    if document.tables[0]:

                        def add_row(table, date_string, text_string):
                            style = document.styles["Normal"]
                            style.paragraph_format.space_before = Pt(3)
                            style.paragraph_format.space_after = Pt(3)

                            row = document.tables[0].add_row()
                            date_paragraph = row.cells[0].paragraphs[0]
                            date_paragraph.text = date_string
                            date_paragraph.style = style
                            date_paragraph.alignment = (
                                WD_ALIGN_PARAGRAPH.CENTER
                            )
                            row.cells[
                                0
                            ].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                            text_paragraph = row.cells[1].paragraphs[0]
                            text_paragraph.text = text_string
                            text_paragraph.style = style
                            text_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            row.cells[
                                1
                            ].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                        add_row(
                            document.tables[0],
                            first_notice_event.date.strftime("%d-%m-%Y"),
                            "Fotos do local",
                        )

                        add_row(
                            document.tables[0],
                            first_notice_event.date.strftime("%d-%m-%Y"),
                            "Vistoria realizada",
                        )

                        intimacoes = notice.notice_events.filter(
                            notice_event_type__name="Intimação"
                        ).all()
                        if intimacoes.count() > 0:
                            for intimacao in intimacoes:
                                add_row(
                                    document.tables[0],
                                    intimacao.date.strftime("%d-%m-%Y"),
                                    "Auto de Intimação n°"
                                    + intimacao.identification,
                                )

                        infracoes = notice.notice_events.filter(
                            notice_event_type__name="Infração"
                        ).all()
                        if infracoes.count() > 0:
                            for infracao in infracoes:
                                add_row(
                                    document.tables[0],
                                    infracao.date.strftime("%d-%m-%Y"),
                                    "Auto de Infração n°"
                                    + infracao.identification,
                                )

                        embargos = notice.notice_events.filter(
                            notice_event_type__name="Embargo"
                        ).all()
                        if embargos.count() > 0:
                            for embargo in embargos:
                                add_row(
                                    document.tables[0],
                                    embargo.date.strftime("%d-%m-%Y"),
                                    "Auto de Embargo n°"
                                    + embargo.identification,
                                )

                        add_row(
                            document.tables[0],
                            timezone.localtime(timezone.now()).strftime(
                                "%d-%m-%Y"
                            ),
                            "Extrato do cadastro de imóvel",
                        )

                        add_row(
                            document.tables[0],
                            timezone.localtime(timezone.now()).strftime(
                                "%d-%m-%Y"
                            ),
                            "Planta quadra",
                        )

                        add_row(
                            document.tables[0],
                            timezone.localtime(timezone.now()).strftime(
                                "%d-%m-%Y"
                            ),
                            "Relatório de fiscalização nº "
                            + report_number
                            + " – encaminhamento para abertura do processo de vistoria administrativa",
                        )

                else:
                    document = Document()
                    document.add_heading("Document Title", 0)

                response = HttpResponse(
                    content_type=(
                        "application/"
                        "vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                )
                response[
                    "Content-Disposition"
                ] = "attachment; filename=download.docx"
                document.save(response)

                return response
        else:
            raise serializers.ValidationError(
                "vistoria_administrativa_id field required."
            )
        return Response(status=status.HTTP_400_BAD_REQUEST)


class FileVARequestDocx(generics.ListCreateAPIView):
    permission_classes = [
        permissions.IsAdminUser,
    ]

    def get(self, request, *args, **kwargs):
        file_path = os.path.join(
            settings.MEDIA_ROOT, "relatorio_padrao", "va_padrao.docx"
        )
        if os.path.exists(file_path):
            document = Document(file_path)
            response = HttpResponse(
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            response[
                "Content-Disposition"
            ] = "attachment; filename=download.docx"
            document.save(response)
            return response
        raise Http404

    def post(self, request, *args, **kwargs):
        dest_dir = file_path = os.path.join(
            settings.MEDIA_ROOT, "relatorio_padrao"
        )
        file_path = os.path.join(dest_dir, "va_padrao.docx")
        file_path_backup = os.path.join(dest_dir, "va_padrao_backup.docx")
        if os.path.exists(file_path):
            if os.path.exists(file_path_backup):
                os.remove(file_path_backup)
            os.rename(file_path, file_path_backup)
        with open(file_path, "wb+") as destination:
            for chunk in request.FILES["va_padrao"].chunks():
                destination.write(chunk)
            return Response(
                {"detail": "Arquivo salvo com sucesso"},
                status=status.HTTP_200_OK,
            )
        return Response(status=status.HTTP_400_BAD_REQUEST)


class FileRFRequestDocx(generics.ListCreateAPIView):
    permission_classes = [
        permissions.IsAdminUser,
    ]

    def get(self, request, *args, **kwargs):
        file_path = os.path.join(
            settings.MEDIA_ROOT, "relatorio_padrao", "rf_padrao.docx"
        )
        if os.path.exists(file_path):
            document = Document(file_path)
            response = HttpResponse(
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            response[
                "Content-Disposition"
            ] = "attachment; filename=download.docx"
            document.save(response)
            return response
        raise Http404

    def post(self, request, *args, **kwargs):
        dest_dir = file_path = os.path.join(
            settings.MEDIA_ROOT, "relatorio_padrao"
        )
        file_path = os.path.join(dest_dir, "rf_padrao.docx")
        file_path_backup = os.path.join(dest_dir, "rf_padrao_backup.docx")
        if os.path.exists(file_path):
            if os.path.exists(file_path_backup):
                os.remove(file_path_backup)
            os.rename(file_path, file_path_backup)
        with open(file_path, "wb+") as destination:
            for chunk in request.FILES["rf_padrao"].chunks():
                destination.write(chunk)
            return Response(
                {"detail": "Arquivo salvo com sucesso"},
                status=status.HTTP_200_OK,
            )
        return Response(status=status.HTTP_400_BAD_REQUEST)


class downloadNotification(generics.RetrieveAPIView):
    permission_classes = [
        permissions.IsAuthenticated,
    ]

    def get(self, request, *args, **kwargs):
        # locale.setlocale(locale.LC_TIME, "pt_BR")
        locale.setlocale(locale.LC_TIME, "pt_BR.utf8")
        notice_event_with_date = parse_bool(self.request.query_params.get(
            "notice_event_with_date", False
        )
        )
        notice_event_reference = self.request.query_params.get(
            "notice_event_reference", None
        )

        notice_event_type_file_id = self.request.query_params.get(
            "notice_event_type_file_id", None
        )

        notice_event_type_file = None
        if notice_event_type_file_id:
            notice_event_type_file = NoticeEventTypeFile.objects.filter(
                id=notice_event_type_file_id
            ).first()

        notice_event_id = self.request.query_params.get(
            "notice_event_id", None
        )
        notice_event = None
        if notice_event_id:
            notice_event = NoticeEvent.objects.filter(
                id=notice_event_id
            ).first()

        if notice_event_type_file and notice_event:

            file_path = notice_event_type_file.file_doc.path

            user = notice_event.notice.owner

            if not self.request.user.profile.is_auditor():
                if (
                    not user.profile.has_my_permission(request.user)
                    or (self.request.user.profile.is_particular() and user != self.request.user)
                ):
                    return Response(status=status.HTTP_403_FORBIDDEN)

            address_string = ""
            if notice_event.notice.imovel.logradouro:
                address_string = notice_event.notice.imovel.logradouro
            if notice_event.notice.imovel.numero:
                address_string += ", n" + notice_event.notice.imovel.numero
            if notice_event.notice.imovel.complemento:
                address_string += ", " + notice_event.notice.imovel.complemento
            if notice_event.notice.imovel.bairro:
                address_string += " - " + notice_event.notice.imovel.bairro

            def toUpperCasoOrNone(string):
                if string:
                    return str(string).upper()
                return None

            now = timezone.localtime(timezone.now())

            data_atual = ""
            hora_atual = ""
            if notice_event_with_date:
                data_atual = now.strftime("%d/%m/%Y")
                hora_atual = now.strftime("%H:%M")

            context = {
                "data_atual": data_atual,
                "hora_atual": hora_atual,
                "data_atual_por_extenso": now.strftime("%d de %B de %Y"),
                "afm_nome_completo": toUpperCasoOrNone(user.get_full_name()),
                "afm_matricula": toUpperCasoOrNone(user.profile.matricula),
                "auto_identificacao": toUpperCasoOrNone(notice_event.identification),
                "auto_data": notice_event.date.strftime("%d/%m/%Y"),
                "auto_data_por_extenso": notice_event.date.strftime("%d de %B de %Y"),
                "auto_documento": toUpperCasoOrNone(notice_event.notice.document),
                "auto_tipo": toUpperCasoOrNone(notice_event.notice_event_type),
                "imovel_razao_social": toUpperCasoOrNone(notice_event.notice.imovel.razao_social),
                "imovel_inscricao": toUpperCasoOrNone(notice_event.notice.imovel.inscricao_imobiliaria),
                "imovel_endereco_completo": toUpperCasoOrNone(address_string),
                "auto_descumprido": toUpperCasoOrNone(notice_event_reference),
            }
            document = docxFromTemplate(file_path, context)
            if document:
                response = HttpResponse(
                    content_type=(
                        "application/"
                        "vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                )
                response[
                    "Content-Disposition"
                ] = "attachment; filename=download.docx"
                document.save(response)
                return response

        return Response(status=status.HTTP_400_BAD_REQUEST)
